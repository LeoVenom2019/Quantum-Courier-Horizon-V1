from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path("D:/PROJETOS/QCH/Curriculo_Leonardo_Dias.docx")

NAVY = "17324D"
BLUE = "2E74B5"
LIGHT_BLUE = "EAF2FB"
PALE = "F6F8FB"
MUTED = "64748B"
INK = "172033"
WHITE = "FFFFFF"
BORDER = "D7DEE8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def remove_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_margin(cell, top=90, start=140, bottom=90, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, color=INK, bold=None, italic=None, name="Aptos"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_text(paragraph, text, size=10.5, color=INK, bold=False, italic=False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return run


def set_para(paragraph, before=0, after=4, line=1.05, align=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_section_title(container, title):
    p = container.add_paragraph()
    set_para(p, before=8, after=5, line=1.0)
    add_text(p, title.upper(), size=10, color=BLUE, bold=True)
    return p


def add_bullet(container, text):
    p = container.add_paragraph(style="List Bullet")
    set_para(p, before=0, after=3, line=1.06)
    add_text(p, text, size=9.7, color=INK)


def add_sidebar_kv(container, label, value):
    p = container.add_paragraph()
    set_para(p, before=0, after=5, line=1.05)
    add_text(p, f"{label}: ", size=8.7, color=MUTED, bold=True)
    add_text(p, value, size=9.2, color=INK)


def add_job(container, company, period, role):
    p = container.add_paragraph()
    set_para(p, before=5, after=1, line=1.0)
    add_text(p, company, size=11.2, color=INK, bold=True)

    p = container.add_paragraph()
    set_para(p, before=0, after=1, line=1.0)
    add_text(p, f"Período: {period}", size=9.2, color=MUTED, bold=True)

    p = container.add_paragraph()
    set_para(p, before=0, after=4, line=1.0)
    add_text(p, role, size=10.2, color=INK)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    bullet = styles["List Bullet"]
    bullet.font.name = "Aptos"
    bullet._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    bullet.paragraph_format.left_indent = Cm(0.45)
    bullet.paragraph_format.first_line_indent = Cm(-0.18)
    bullet.paragraph_format.space_after = Pt(3)

    # Header block with photo placeholder.
    header_table = doc.add_table(rows=1, cols=2)
    set_table_width(header_table, [6900, 1500])
    for cell in header_table.rows[0].cells:
        remove_cell_border(cell)
        set_cell_margin(cell, 130, 180, 130, 180)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left, right = header_table.rows[0].cells
    set_cell_shading(left, NAVY)
    set_cell_shading(right, NAVY)

    p = left.paragraphs[0]
    set_para(p, after=1, line=1.0)
    add_text(p, "LEONARDO DIAS", size=25, color=WHITE, bold=True)
    p = left.add_paragraph()
    set_para(p, after=0, line=1.0)
    add_text(p, "Currículo Profissional", size=11.5, color="DDEAF7", bold=True)

    photo = right.paragraphs[0]
    photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(photo, after=0, line=1.0)
    add_text(photo, "\nESPAÇO\nPARA\nFOTO\n", size=8.8, color=WHITE, bold=True)
    set_cell_border(right, color="9AB3CF", size="12")

    spacer = doc.add_paragraph()
    set_para(spacer, after=6)

    layout = doc.add_table(rows=1, cols=2)
    set_table_width(layout, [5600, 2800])
    layout.autofit = False
    main_cell, side_cell = layout.rows[0].cells
    for cell in (main_cell, side_cell):
        remove_cell_border(cell)
        set_cell_margin(cell, 40, 120, 40, 120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_shading(side_cell, PALE)
    set_cell_border(side_cell, color=BORDER, size="8")

    add_section_title(main_cell, "Experiência de Trabalho")
    add_job(main_cell, "Capixaba Comércio de Produtos", "05/01/2021 - 01/04/2025", "Frentista (Caixa)")
    add_job(main_cell, "Guarave Guarapari Veículos LTDA", "10/06/2013 - 02/02/2020", "Frentista (Caixa)")
    add_job(main_cell, "Geração Speed Kart LTDA", "22/12/2011 - 04/05/2013", "Motorista de Carro de Passeio")
    add_job(main_cell, "RG Comércio de Vidros LTDA", "02/08/2007 - 31/05/2008", "Vidraceiro")

    add_section_title(side_cell, "Dados Pessoais")
    add_sidebar_kv(side_cell, "Nascimento", "25/06/1982")
    add_sidebar_kv(side_cell, "Naturalidade", "Águas Formosas - MG")
    add_sidebar_kv(side_cell, "Residência", "Guarapari - ES")
    add_sidebar_kv(side_cell, "Desde", "1990")

    add_section_title(side_cell, "Educação")
    add_bullet(side_cell, "Ensino Médio Completo")
    add_bullet(side_cell, "Ensino Superior Incompleto")

    add_section_title(side_cell, "Habilidades")
    for item in [
        "Ótimo convívio com colegas",
        "Proatividade e iniciativa",
        "Honestidade",
        "Disponibilidade de horário",
        "Carteira de motorista AB",
        "Não fumante",
    ]:
        add_bullet(side_cell, item)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(footer, after=0, line=1.0)
    add_text(footer, "Leonardo Dias", size=8.5, color=MUTED)

    doc.core_properties.title = "Currículo - Leonardo Dias"
    doc.core_properties.author = "Leonardo Dias"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
